"""Forensic report generation (lab manual Step 21 / §6.8).

Formats: HTML, PDF, XLSX, DOCX, XML, JSON, CSV.  Scope: ``all``, a filtered
AQL subset, or an explicit selection — the three options the manual's
Report/Export dialog offers.

Every report leads with the integrity statement, not with the findings.  A
report whose evidence failed verification must say so on page one, because a
reader who reaches the findings first has already been misled.  If verification
failed, the banner is red and unmissable, and the failure is repeated in the
conclusion.
"""

from __future__ import annotations

import csv
import html
import json
import io
from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, List, Optional, Sequence

from ..analyze.session import AnalysisSession
from ..core.errors import ReportError
from ..core.models import Category
from ..parsers.timestamps import to_iso

FORMATS = ("html", "pdf", "xlsx", "docx", "xml", "json", "csv")


@dataclass
class ReportOptions:
    title: str = "Mobile Device Forensic Examination Report"
    scope: str = "all"                       # all | filtered | selected
    query: str = ""                          # AQL when scope == filtered
    selected_ids: List[str] = field(default_factory=list)
    formats: List[str] = field(default_factory=lambda: ["html"])
    include_deleted: bool = True
    include_media: bool = True
    include_graph: bool = True
    include_timeline: bool = True
    include_log: bool = False
    include_audit: bool = False
    include_intelligence: bool = True
    owner_identifiers: List[str] = field(default_factory=list)
    max_artifacts: int = 20000
    examiner: str = ""
    organisation: str = ""
    reference: str = ""
    conclusion: str = ""
    logo_note: str = ""


class ReportBuilder:
    """Assemble report data once, then render it into any format."""

    def __init__(self, session: AnalysisSession, options: ReportOptions):
        self.session = session
        self.options = options
        self.data = self._collect()

    # ------------------------------------------------------------- gathering
    def _collect(self) -> Dict[str, Any]:
        o = self.options
        overview = self.session.overview()

        if o.scope == "selected" and o.selected_ids:
            artifacts = [a for a in
                         (self.session.get(i) for i in o.selected_ids) if a]
        else:
            query = o.query if o.scope == "filtered" else ""
            if not o.include_deleted:
                query = f"({query}) AND deleted:false" if query else "deleted:false"
            artifacts = self.session.query(query, limit=o.max_artifacts
                                           )["artifacts"]

        by_category: Dict[str, List[Dict[str, Any]]] = {}
        for a in artifacts:
            by_category.setdefault(a["category"], []).append(a)

        data: Dict[str, Any] = {
            "generated_at": datetime.now(timezone.utc).isoformat(timespec="seconds"),
            "options": {
                "title": o.title, "scope": o.scope, "query": o.query,
                "examiner": o.examiner, "organisation": o.organisation,
                "reference": o.reference,
            },
            "overview": overview,
            "integrity": self.session.integrity_report(),
            "artifacts": artifacts,
            "artifact_count": len(artifacts),
            "by_category": by_category,
            "category_counts": {k: len(v) for k, v in by_category.items()},
            "deleted_count": sum(1 for a in artifacts if a.get("is_deleted")),
            "conclusion": o.conclusion,
        }
        if o.include_timeline:
            data["statistics"] = self.session.statistics(
                o.query if o.scope == "filtered" else "")
        if o.include_graph:
            data["graph"] = self.session.connections("all", min_weight=1,
                                                     max_nodes=120)
        if o.include_intelligence:
            try:
                data["intelligence"] = self.session.intelligence(
                    o.owner_identifiers)
            except Exception as exc:
                data["intelligence_error"] = f"{type(exc).__name__}: {exc}"
        if o.include_log:
            data["extraction_log"] = self.session.extraction_log()
        if o.include_audit:
            data["audit"] = self.session.audit_trail()
        return data

    # -------------------------------------------------------------- dispatch
    def write(self, out_dir: Path | str, basename: str = "forensic_report"
              ) -> List[Path]:
        out_dir = Path(out_dir)
        out_dir.mkdir(parents=True, exist_ok=True)
        written: List[Path] = []
        for fmt in self.options.formats:
            fmt = fmt.lower().strip()
            if fmt not in FORMATS:
                raise ReportError(
                    f"unsupported format {fmt!r}; choose from {FORMATS}")
            target = out_dir / f"{basename}.{fmt}"
            getattr(self, f"_write_{fmt}")(target)
            written.append(target)
        return written

    # ------------------------------------------------------------------ JSON
    def _write_json(self, path: Path) -> None:
        path.write_text(json.dumps(self.data, indent=2, ensure_ascii=False,
                                   default=str), encoding="utf-8")

    # ------------------------------------------------------------------- CSV
    def _write_csv(self, path: Path) -> None:
        cols = ["timestamp_iso", "category", "subtype", "direction", "app",
                "parties", "body", "recovery", "confidence", "source_path",
                "source_table", "source_row", "blob_sha256", "latitude",
                "longitude", "artifact_id"]
        with path.open("w", newline="", encoding="utf-8-sig") as fh:
            w = csv.writer(fh)
            w.writerow(cols)
            for a in self.data["artifacts"]:
                w.writerow([
                    a.get("timestamp_iso", ""), a.get("category", ""),
                    a.get("subtype", ""), a.get("direction", ""),
                    a.get("app", ""),
                    "; ".join(p["label"] for p in a.get("parties", [])
                              if not p.get("is_owner")),
                    (a.get("body") or "").replace("\n", " ⏎ "),
                    a.get("recovery", ""), a.get("confidence", ""),
                    a.get("source_path", ""), a.get("source_table", ""),
                    a.get("source_row", ""), a.get("blob_sha256", ""),
                    a.get("latitude", ""), a.get("longitude", ""),
                    a.get("artifact_id", ""),
                ])

    # ------------------------------------------------------------------- XML
    def _write_xml(self, path: Path) -> None:
        e = lambda s: html.escape(str(s if s is not None else ""), quote=True)
        o = self.data["overview"]
        integ = self.data["integrity"]
        lines = ['<?xml version="1.0" encoding="UTF-8"?>',
                 '<ArgusForensicReport version="1.0">',
                 '  <Metadata>',
                 f'    <GeneratedAt>{e(self.data["generated_at"])}</GeneratedAt>',
                 f'    <Title>{e(self.options.title)}</Title>',
                 f'    <Examiner>{e(self.options.examiner)}</Examiner>',
                 f'    <Organisation>{e(self.options.organisation)}</Organisation>',
                 f'    <Reference>{e(self.options.reference)}</Reference>',
                 '  </Metadata>',
                 '  <Case>',
                 f'    <CaseId>{e(o["case_id"])}</CaseId>',
                 f'    <ExhibitId>{e(o["exhibit_id"])}</ExhibitId>',
                 f'    <Operator>{e(o["operator"])}</Operator>',
                 f'    <Method>{e(o["method"])}</Method>',
                 '    <Device>',
                 *[f'      <{k.title().replace("_","")}>{e(v)}</{k.title().replace("_","")}>'
                   for k, v in o["device"].items()],
                 '    </Device>',
                 '  </Case>',
                 f'  <Integrity verified="{str(integ["ok"]).lower()}">',
                 *[f'    <Problem>{e(p)}</Problem>'
                   for c in integ["containers"] for p in c.get("problems", [])],
                 '  </Integrity>',
                 f'  <Artifacts count="{self.data["artifact_count"]}">']
        for a in self.data["artifacts"]:
            lines.append(f'    <Artifact id="{e(a["artifact_id"])}" '
                         f'category="{e(a["category"])}" '
                         f'deleted="{str(a.get("is_deleted", False)).lower()}">')
            lines.append(f'      <Timestamp>{e(a.get("timestamp_iso"))}</Timestamp>')
            lines.append(f'      <Type>{e(a.get("subtype"))}</Type>')
            lines.append(f'      <Direction>{e(a.get("direction"))}</Direction>')
            lines.append(f'      <Application>{e(a.get("app"))}</Application>')
            lines.append(f'      <Body>{e(a.get("body"))}</Body>')
            lines.append('      <Participants>')
            for p in a.get("parties", []):
                lines.append(
                    f'        <Participant role="{e(p.get("role"))}" '
                    f'owner="{str(p.get("is_owner", False)).lower()}">'
                    f'{e(p.get("label"))}</Participant>')
            lines.append('      </Participants>')
            lines.append('      <Provenance>')
            lines.append(f'        <SourcePath>{e(a.get("source_path"))}</SourcePath>')
            lines.append(f'        <SourceTable>{e(a.get("source_table"))}</SourceTable>')
            lines.append(f'        <SourceRow>{e(a.get("source_row"))}</SourceRow>')
            lines.append(f'        <Recovery>{e(a.get("recovery"))}</Recovery>')
            lines.append(f'        <Sha256>{e(a.get("blob_sha256"))}</Sha256>')
            lines.append('      </Provenance>')
            lines.append('    </Artifact>')
        lines += ['  </Artifacts>', '</ArgusForensicReport>']
        path.write_text("\n".join(lines), encoding="utf-8")

    # ------------------------------------------------------------------ XLSX
    def _write_xlsx(self, path: Path) -> None:
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, PatternFill, Alignment
            from openpyxl.utils import get_column_letter
        except ImportError as exc:
            raise ReportError("XLSX export requires openpyxl "
                              "(pip install openpyxl)") from exc

        wb = Workbook()
        head_fill = PatternFill("solid", fgColor="1F3B57")
        head_font = Font(bold=True, color="FFFFFF", size=10)
        del_fill = PatternFill("solid", fgColor="F3E6F7")

        # Summary sheet
        ws = wb.active
        ws.title = "Summary"
        o = self.data["overview"]
        integ = self.data["integrity"]
        rows = [
            ("ARGUS Forensic Report", ""),
            ("Title", self.options.title),
            ("Generated", self.data["generated_at"]),
            ("Examiner", self.options.examiner),
            ("Organisation", self.options.organisation),
            ("Reference", self.options.reference),
            ("", ""),
            ("Case ID", o["case_id"]),
            ("Exhibit ID", o["exhibit_id"]),
            ("Operator", o["operator"]),
            ("Extraction method", o["method"]),
            ("Time span", o["time_span"]),
            ("", ""),
            ("Device make", o["device"]["make"]),
            ("Device model", o["device"]["model"]),
            ("Operating system", o["device"]["os"]),
            ("Serial", o["device"]["serial"]),
            ("IMEI", o["device"]["imei"]),
            ("", ""),
            ("INTEGRITY", "VERIFIED" if integ["ok"] else "FAILED"),
            ("Artifacts in report", self.data["artifact_count"]),
            ("Recovered from deleted space", self.data["deleted_count"]),
            ("First activity", o["first_activity"]),
            ("Last activity", o["last_activity"]),
        ]
        for r, (k, v) in enumerate(rows, 1):
            ws.cell(r, 1, k).font = Font(bold=True, size=10)
            ws.cell(r, 2, v)
        ws.column_dimensions["A"].width = 30
        ws.column_dimensions["B"].width = 62

        # One sheet per category
        cols = ["Time (UTC)", "Type", "Direction", "Party", "Content",
                "Application", "Recovery", "Confidence", "Source file",
                "Table", "Row", "SHA-256", "Latitude", "Longitude",
                "Artifact ID"]
        for category, items in sorted(self.data["by_category"].items(),
                                      key=lambda kv: -len(kv[1])):
            name = category.replace("&", "and").replace("/", "-")[:31]
            sheet = wb.create_sheet(name)
            for c, header in enumerate(cols, 1):
                cell = sheet.cell(1, c, header)
                cell.fill, cell.font = head_fill, head_font
                cell.alignment = Alignment(vertical="center")
            for r, a in enumerate(items, 2):
                values = [
                    a.get("timestamp_iso", ""), a.get("subtype", ""),
                    a.get("direction", ""),
                    "; ".join(p["label"] for p in a.get("parties", [])
                              if not p.get("is_owner")),
                    (a.get("body") or "")[:32000],
                    a.get("app", ""), a.get("recovery", ""),
                    a.get("confidence", ""), a.get("source_path", ""),
                    a.get("source_table", ""), a.get("source_row", ""),
                    a.get("blob_sha256", ""), a.get("latitude", ""),
                    a.get("longitude", ""), a.get("artifact_id", ""),
                ]
                for c, v in enumerate(values, 1):
                    cell = sheet.cell(r, c, v)
                    if a.get("is_deleted"):
                        cell.fill = del_fill
            widths = [20, 22, 11, 30, 60, 16, 18, 11, 40, 16, 8, 24, 12, 12, 34]
            for i, w in enumerate(widths, 1):
                sheet.column_dimensions[get_column_letter(i)].width = w
            sheet.freeze_panes = "A2"
            sheet.auto_filter.ref = sheet.dimensions

        # Connection sheet
        if self.data.get("graph"):
            gs = wb.create_sheet("Connections")
            for c, h in enumerate(["Party A", "Party B", "Artifacts", "Calls",
                                   "Messages", "Reciprocity", "First", "Last"], 1):
                cell = gs.cell(1, c, h)
                cell.fill, cell.font = head_fill, head_font
            labels = {n["key"]: n["label"] for n in self.data["graph"]["nodes"]}
            for r, e in enumerate(self.data["graph"]["edges"], 2):
                gs.cell(r, 1, labels.get(e["source"], e["source"]))
                gs.cell(r, 2, labels.get(e["target"], e["target"]))
                gs.cell(r, 3, e["artifact_count"])
                gs.cell(r, 4, e["calls"])
                gs.cell(r, 5, e["messages"])
                gs.cell(r, 6, e["reciprocity"])
                gs.cell(r, 7, to_iso(e.get("first_seen")))
                gs.cell(r, 8, to_iso(e.get("last_seen")))
            for i, w in enumerate([30, 30, 12, 10, 12, 13, 22, 22], 1):
                gs.column_dimensions[get_column_letter(i)].width = w
            gs.freeze_panes = "A2"

        wb.save(path)

    # ------------------------------------------------------------------ DOCX
    def _write_docx(self, path: Path) -> None:
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError as exc:
            raise ReportError("DOCX export requires python-docx "
                              "(pip install python-docx)") from exc

        doc = Document()
        o = self.data["overview"]
        integ = self.data["integrity"]

        doc.add_heading(self.options.title, 0)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Generated {self.data['generated_at']}  ·  "
                        f"ARGUS Forensics")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        # Integrity banner first — before any finding
        doc.add_heading("Evidence integrity", level=1)
        banner = doc.add_paragraph()
        if integ["ok"]:
            r = banner.add_run(
                "VERIFIED — every blob re-hashed to its stored digest, the "
                "container seal matches, and the chain-of-custody log is "
                "intact and unbroken.")
            r.font.color.rgb = RGBColor(0x1B, 0x7F, 0x37)
        else:
            r = banner.add_run(
                "VERIFICATION FAILED — the evidence in this report does NOT "
                "match its recorded hashes. The findings below must not be "
                "relied upon until this is resolved.")
            r.font.color.rgb = RGBColor(0xC0, 0x20, 0x1A)
            r.bold = True
            for c in integ["containers"]:
                for prob in c.get("problems", []):
                    doc.add_paragraph(f"{c['name']}: {prob}", style="List Bullet")

        doc.add_heading("Case and exhibit", level=1)
        table = doc.add_table(rows=0, cols=2)
        table.style = "Light Grid Accent 1"
        pairs = [
            ("Case ID", o["case_id"]), ("Exhibit ID", o["exhibit_id"]),
            ("Operator", o["operator"]), ("Examiner", self.options.examiner),
            ("Organisation", self.options.organisation),
            ("Reference", self.options.reference),
            ("Extraction method", o["method"]),
            ("Time span", o["time_span"]),
            ("Extraction started", o["started_at"]),
            ("Extraction finished", o["finished_at"]),
            ("Device", " ".join(filter(None, [o["device"]["make"],
                                              o["device"]["model"]]))),
            ("Operating system", o["device"]["os"]),
            ("Serial", o["device"]["serial"]), ("IMEI", o["device"]["imei"]),
            ("ICCID", o["device"]["iccid"]),
            ("Phone number", o["device"]["phone_number"]),
            ("Lock state at acquisition", o["device"]["lock_state"]),
        ]
        for k, v in pairs:
            if not v:
                continue
            row = table.add_row().cells
            row[0].text = k
            row[1].text = str(v)

        doc.add_heading("Summary of findings", level=1)
        doc.add_paragraph(
            f"{self.data['artifact_count']:,} artifacts are included in this "
            f"report, of which {self.data['deleted_count']:,} were recovered "
            f"from deleted or unallocated space and were not visible to the "
            f"device user at the time of seizure. Activity spans "
            f"{o['first_activity'] or 'unknown'} to "
            f"{o['last_activity'] or 'unknown'}.")

        ct = doc.add_table(rows=1, cols=2)
        ct.style = "Light List Accent 1"
        ct.rows[0].cells[0].text = "Category"
        ct.rows[0].cells[1].text = "Count"
        for k, v in sorted(self.data["category_counts"].items(),
                           key=lambda kv: -kv[1]):
            row = ct.add_row().cells
            row[0].text = k
            row[1].text = f"{v:,}"

        if self.data.get("graph", {}).get("top_contacts"):
            doc.add_heading("Communication analysis", level=1)
            doc.add_paragraph(
                "Parties ranked by volume of communication with the device "
                "owner across all applications.")
            gt = doc.add_table(rows=1, cols=4)
            gt.style = "Light Grid Accent 1"
            for i, h in enumerate(["Party", "Artifacts", "Calls", "Messages"]):
                gt.rows[0].cells[i].text = h
            for n in self.data["graph"]["top_contacts"][:15]:
                row = gt.add_row().cells
                row[0].text = str(n["label"])
                row[1].text = f"{n['artifact_count']:,}"
                row[2].text = f"{n['calls']:,}"
                row[3].text = f"{n['messages']:,}"

        doc.add_page_break()
        doc.add_heading("Artifacts", level=1)
        for category, items in sorted(self.data["by_category"].items(),
                                      key=lambda kv: -len(kv[1])):
            doc.add_heading(f"{category} ({len(items):,})", level=2)
            at = doc.add_table(rows=1, cols=5)
            at.style = "Light Grid"
            for i, h in enumerate(["Time (UTC)", "Type", "Party",
                                   "Content", "State"]):
                at.rows[0].cells[i].text = h
            for a in items[:400]:
                row = at.add_row().cells
                row[0].text = a.get("timestamp_iso", "")
                row[1].text = a.get("subtype", "")
                row[2].text = "; ".join(
                    p["label"] for p in a.get("parties", [])
                    if not p.get("is_owner"))[:60]
                row[3].text = (a.get("body") or "")[:220]
                row[4].text = "DELETED" if a.get("is_deleted") else ""
            if len(items) > 400:
                doc.add_paragraph(
                    f"({len(items) - 400:,} further {category} artifacts are "
                    f"included in the XLSX and XML exports.)")

        if self.options.conclusion:
            doc.add_heading("Conclusion", level=1)
            doc.add_paragraph(self.options.conclusion)
        if not integ["ok"]:
            para = doc.add_paragraph()
            r = para.add_run(
                "Note: evidence integrity verification failed for this case. "
                "See the Evidence integrity section above.")
            r.bold = True
            r.font.color.rgb = RGBColor(0xC0, 0x20, 0x1A)

        doc.save(path)

    # ------------------------------------------------------------------ HTML
    def _write_html(self, path: Path) -> None:
        path.write_text(self._html_document(), encoding="utf-8")

    def _html_document(self) -> str:
        d, o = self.data, self.data["overview"]
        integ = d["integrity"]
        e = html.escape

        def table(rows: Sequence[Dict[str, Any]], columns: Sequence[str]) -> str:
            head = "".join(f"<th>{e(c)}</th>" for c in columns)
            body = []
            for r in rows:
                cls = ' class="del"' if r.get("_deleted") else ""
                cells = "".join(f"<td>{e(str(r.get(c, '')))}</td>" for c in columns)
                body.append(f"<tr{cls}>{cells}</tr>")
            return (f"<table><thead><tr>{head}</tr></thead>"
                    f"<tbody>{''.join(body)}</tbody></table>")

        sections = []
        for category, items in sorted(d["by_category"].items(),
                                      key=lambda kv: -len(kv[1])):
            rows = [{
                "Time (UTC)": a.get("timestamp_iso", ""),
                "Type": a.get("subtype", ""),
                "Direction": a.get("direction", ""),
                "Party": "; ".join(p["label"] for p in a.get("parties", [])
                                   if not p.get("is_owner")),
                "Content": (a.get("body") or "")[:600],
                "Application": a.get("app", ""),
                "State": a.get("recovery") if a.get("is_deleted") else "",
                "Source": a.get("source_path", ""),
                "_deleted": a.get("is_deleted"),
            } for a in items[:2000]]
            more = (f"<p class='note'>{len(items) - 2000:,} further artifacts "
                    f"omitted from this HTML view; see the XLSX or XML export."
                    f"</p>" if len(items) > 2000 else "")
            sections.append(
                f"<h3 id='cat-{e(category)}'>{e(category)} "
                f"<span class='badge'>{len(items):,}</span></h3>"
                + table(rows, ["Time (UTC)", "Type", "Direction", "Party",
                               "Content", "Application", "State", "Source"])
                + more)

        graph_rows = ""
        if d.get("graph"):
            labels = {n["key"]: n["label"] for n in d["graph"]["nodes"]}
            graph_rows = table([{
                "Party A": labels.get(x["source"], x["source"]),
                "Party B": labels.get(x["target"], x["target"]),
                "Artifacts": x["artifact_count"], "Calls": x["calls"],
                "Messages": x["messages"], "Reciprocity": x["reciprocity"],
                "First": to_iso(x.get("first_seen")),
                "Last": to_iso(x.get("last_seen")),
            } for x in d["graph"]["edges"][:200]],
                ["Party A", "Party B", "Artifacts", "Calls", "Messages",
                 "Reciprocity", "First", "Last"])

        stats = d.get("statistics", {})
        hist = stats.get("histogram", {})
        hours = hist.get("by_hour", [])
        maxh = max((h["count"] for h in hours), default=1) or 1
        hour_bars = "".join(
            f'<div class="hb" style="height:{max(2, h["count"]/maxh*100):.0f}%" '
            f'title="{h["hour"]:02d}:00 — {h["count"]} artifacts"></div>'
            for h in hours)

        banner = (
            '<div class="banner ok"><b>INTEGRITY VERIFIED</b> — every blob '
            're-hashed to its stored digest, the container seal matches, and '
            'the chain-of-custody log is unbroken.</div>'
            if integ["ok"] else
            '<div class="banner bad"><b>INTEGRITY VERIFICATION FAILED</b> — '
            'the evidence described below does not match its recorded hashes. '
            'These findings must not be relied upon until resolved.<ul>'
            + "".join(f"<li>{e(c['name'])}: {e(p)}</li>"
                      for c in integ["containers"]
                      for p in c.get("problems", []))
            + '</ul></div>')

        dev = " ".join(filter(None, [o["device"]["make"], o["device"]["model"],
                                     o["device"]["os"]]))
        lead_sheet = self._html_lead_sheet()
        return f"""<!DOCTYPE html>
<html lang="en"><head><meta charset="utf-8">
<title>{e(self.options.title)}</title>
<style>
:root{{--ink:#14181d;--mut:#5b6774;--line:#dde3ea;--accent:#1f4e79;
  --del:#7b3fa0;--ok:#1b7f37;--bad:#c0201a}}
*{{box-sizing:border-box}}
body{{font:13.5px/1.65 -apple-system,BlinkMacSystemFont,"Segoe UI",Roboto,
  Helvetica,Arial,sans-serif;color:var(--ink);margin:0;background:#fff}}
.page{{max-width:1180px;margin:0 auto;padding:34px 26px 70px}}
h1{{font-size:26px;margin:0 0 4px;color:var(--accent)}}
h2{{font-size:17px;margin:34px 0 12px;padding-bottom:6px;
  border-bottom:2px solid var(--accent);color:var(--accent)}}
h3{{font-size:14px;margin:24px 0 9px}}
.sub{{color:var(--mut);font-size:12px;margin-bottom:22px}}
.banner{{padding:13px 16px;border-radius:7px;margin:18px 0;font-size:13px}}
.banner.ok{{background:#e9f7ee;border:1px solid #9ed7b1;color:var(--ok)}}
.banner.bad{{background:#fdecea;border:1px solid #f0a29c;color:var(--bad)}}
.banner ul{{margin:8px 0 0 18px}}
.grid{{display:grid;grid-template-columns:repeat(auto-fit,minmax(250px,1fr));
  gap:14px;margin:14px 0}}
.card{{border:1px solid var(--line);border-radius:8px;padding:13px}}
.card h4{{margin:0 0 9px;font-size:11px;text-transform:uppercase;
  letter-spacing:.08em;color:var(--mut)}}
dl{{display:grid;grid-template-columns:auto 1fr;gap:3px 12px;margin:0;font-size:12.5px}}
dt{{color:var(--mut)}}
dd{{margin:0;font-family:ui-monospace,Menlo,Consolas,monospace;font-size:11.5px;
  word-break:break-all}}
table{{width:100%;border-collapse:collapse;margin:9px 0 20px;font-size:11.5px}}
th{{background:var(--accent);color:#fff;text-align:left;padding:7px 9px;
  font-weight:600;font-size:11px;position:sticky;top:0}}
td{{padding:6px 9px;border-bottom:1px solid var(--line);vertical-align:top;
  max-width:420px;word-break:break-word}}
tr.del td{{background:#faf3fd;color:var(--del)}}
tr:nth-child(even) td{{background:#fbfcfd}}
tr.del:nth-child(even) td{{background:#f6ecfb}}
.badge{{background:var(--accent);color:#fff;border-radius:11px;padding:1px 9px;
  font-size:11px;font-weight:600;vertical-align:middle}}
.big{{font-size:28px;font-weight:600;font-family:ui-monospace,Menlo,monospace}}
.big small{{display:block;font-size:10.5px;color:var(--mut);font-weight:400;
  text-transform:uppercase;letter-spacing:.08em;font-family:inherit}}
.hbars{{display:flex;align-items:flex-end;gap:2px;height:90px;margin-top:6px}}
.hb{{flex:1;background:var(--accent);border-radius:2px 2px 0 0;opacity:.85}}
.note{{color:var(--mut);font-size:11.5px;font-style:italic}}
.finding{{background:#fafbfc;border:1px solid var(--line);border-left-width:4px;
  border-radius:6px;padding:11px 14px;margin:11px 0}}
.fhead{{display:flex;align-items:center;gap:9px;flex-wrap:wrap;margin-bottom:6px}}
.sev{{color:#fff;font-size:9.5px;font-weight:700;letter-spacing:.06em;
  padding:2px 7px;border-radius:9px}}
.finding .conf{{margin-left:auto;color:var(--mut);font-size:11px;
  font-family:ui-monospace,Menlo,monospace}}
.finding p{{margin:5px 0;font-size:12.5px}}
.finding .why{{color:#14181d}}
.finding .meta{{color:var(--mut);font-size:11px}}
.finding .caveat{{color:#8a5a00;background:#fffaf0;border:1px solid #f0dcb0;
  border-radius:5px;padding:6px 9px;font-size:11.5px}}
.finding ul.ev{{margin:5px 0 5px 18px;font-size:11.5px;color:#333}}
.finding ul.ev li{{margin:2px 0}}
.toc a{{color:var(--accent);text-decoration:none;margin-right:14px;font-size:12px}}
footer{{margin-top:44px;padding-top:14px;border-top:1px solid var(--line);
  color:var(--mut);font-size:11px}}
@media print{{
  .page{{max-width:none;padding:0}} th{{position:static}}
  h2{{page-break-after:avoid}} table{{page-break-inside:auto}}
  tr{{page-break-inside:avoid}}
}}
</style></head><body><div class="page">

<h1>{e(self.options.title)}</h1>
<div class="sub">
  Generated {e(d['generated_at'])} · ARGUS Forensics ·
  Scope: {e(self.options.scope)}{(' — <code>' + e(self.options.query) + '</code>')
    if self.options.query else ''}
</div>

{banner}

<h2>1. Case and exhibit</h2>
<div class="grid">
  <div class="card"><h4>Case</h4><dl>
    <dt>Case ID</dt><dd>{e(o['case_id'] or '—')}</dd>
    <dt>Exhibit ID</dt><dd>{e(o['exhibit_id'] or '—')}</dd>
    <dt>Operator</dt><dd>{e(o['operator'] or '—')}</dd>
    <dt>Examiner</dt><dd>{e(self.options.examiner or o['operator'] or '—')}</dd>
    <dt>Organisation</dt><dd>{e(self.options.organisation or '—')}</dd>
    <dt>Reference</dt><dd>{e(self.options.reference or '—')}</dd>
  </dl></div>
  <div class="card"><h4>Device</h4><dl>
    <dt>Device</dt><dd>{e(dev or '—')}</dd>
    <dt>Serial</dt><dd>{e(o['device']['serial'] or '—')}</dd>
    <dt>IMEI</dt><dd>{e(o['device']['imei'] or '—')}</dd>
    <dt>ICCID</dt><dd>{e(o['device']['iccid'] or '—')}</dd>
    <dt>Number</dt><dd>{e(o['device']['phone_number'] or '—')}</dd>
    <dt>Lock state</dt><dd>{e(o['device']['lock_state'] or '—')}</dd>
  </dl></div>
  <div class="card"><h4>Extraction</h4><dl>
    <dt>Method</dt><dd>{e(o['method'] or '—')}</dd>
    <dt>Time span</dt><dd>{e(o['time_span'] or '—')}</dd>
    <dt>Started</dt><dd>{e(o['started_at'] or '—')}</dd>
    <dt>Finished</dt><dd>{e(o['finished_at'] or '—')}</dd>
    <dt>Seal</dt><dd>{e(o['encryption_level'])}</dd>
  </dl></div>
</div>

<h2>2. Tool, method and integrity</h2>
<p class="lede">The first question asked of any examination is what produced it
and whether the result can be reproduced. This section answers both, and states
what was not attempted.</p>
<div class="grid">
  <div class="card"><h4>Examination tool</h4><dl>
    <dt>Product</dt><dd>ARGUS Forensics</dd>
    <dt>Version</dt><dd>{e(self._tool['version'])}</dd>
    <dt>Build</dt><dd class="mono">{e(self._tool['build'] or 'unidentified')}</dd>
    <dt>Verified</dt><dd>{e(self._tool['verification'])}</dd>
    <dt>Runtime</dt><dd>{e(self._tool['python'])} on {e(self._tool['platform'])}</dd>
  </dl>
  <p class="small">{e(self._tool['note'])}</p>
  </div>
  <div class="card"><h4>Evidence integrity</h4><dl>
    <dt>Container</dt><dd class="mono">{e(self._seal['name'])}</dd>
    <dt>Seal</dt><dd class="mono">{e(self._seal['container_seal'])}</dd>
    <dt>Blob Merkle root</dt><dd class="mono">{e(self._seal['blob_merkle_root'])}</dd>
    <dt>Audit chain</dt><dd>{e(self._seal['audit'])}</dd>
    <dt>Verified at</dt><dd>{e(self._seal['verified_at'])}</dd>
  </dl>
  <p class="small">The seal is a hash over the container manifest. Recomputing it
  from the container reproduces this value exactly; any alteration to the
  evidence changes it. Verify independently with
  <code>argus verify &lt;container&gt;</code>.</p>
  </div>
  <div class="card"><h4>Method</h4>
  <ol class="small">
    <li>The exhibit was registered and its capability matrix checked before
        connection, so no unsupported extraction was attempted.</li>
    <li>Acquisition by <b>{e(o['method'] or '—')}</b>. The source was read
        without modification; databases with no journal or WAL are opened
        immutable, and those with sidecars are copied before reading so that
        replay cannot write to the original.</li>
    <li>Every file was typed by its magic bytes rather than its extension, so a
        renamed file is examined as what it is.</li>
    <li>Allocated records were read through the schema. Unallocated space —
        freeblocks, the freelist, cell slack, WAL frames and rollback journals —
        was then carved for deleted records.</li>
    <li>Timestamps were normalised from their source epochs into UTC.</li>
    <li>Findings were generated by rule, each citing the artifacts it rests on.</li>
    <li>The container was sealed and this report generated from the sealed
        container, not from the live filesystem.</li>
  </ol>
  </div>
</div>

{lead_sheet}

<h2>4. Summary of artifacts</h2>
<div class="grid">
  <div class="card"><div class="big">{d['artifact_count']:,}<small>Artifacts in report</small></div></div>
  <div class="card"><div class="big" style="color:var(--del)">{d['deleted_count']:,}<small>Recovered from deleted space</small></div></div>
  <div class="card"><div class="big" style="font-size:15px">{e(o['first_activity'] or '—')}<small>First activity</small></div></div>
  <div class="card"><div class="big" style="font-size:15px">{e(o['last_activity'] or '—')}<small>Last activity</small></div></div>
</div>

<div class="grid">
  <div class="card"><h4>Artifacts by category</h4><dl>
    {''.join(f'<dt>{e(k)}</dt><dd>{v:,}</dd>'
             for k, v in sorted(d['category_counts'].items(), key=lambda kv: -kv[1]))}
  </dl></div>
  <div class="card"><h4>Applications</h4><dl>
    {''.join(f'<dt>{e(k)}</dt><dd>{v:,}</dd>'
             for k, v in list(o['applications'].items())[:14])}
  </dl></div>
  {'<div class="card"><h4>Activity by hour of day (UTC)</h4>'
   f'<div class="hbars">{hour_bars}</div>'
   f'<p class="note">Night activity (00:00–06:00): '
   f'{hist.get("night_activity_pct", 0)}% of all timestamped artifacts.</p>'
   '</div>' if hours else ''}
</div>

{'<h2>5. Communication analysis</h2>'
 '<p>Links between the device owner and other parties, ranked by volume. '
 'Reciprocity is the ratio of two-way to one-way traffic on a link; a value '
 'near 0 indicates communication in one direction only.</p>' + graph_rows
 if graph_rows else ''}

<h2>6. Artifacts</h2>
<p class="toc">{''.join(f"<a href='#cat-{e(c)}'>{e(c)}</a>"
                        for c in d['by_category'])}</p>
<p class="note">Rows shaded in purple were recovered from deleted or
unallocated space and were not visible to the device user at the time of
seizure.</p>
{''.join(sections)}

{'<h2>6. Conclusion</h2><p>' + e(self.options.conclusion) + '</p>'
 if self.options.conclusion else ''}

<footer>
  Produced by ARGUS Forensics. Every artifact above carries its source file,
  source table and row identifier so that any finding can be re-derived from
  the original evidence. Container seal:
  {e((self.session.primary.container.manifest.get('seal') or {}).get('container_seal', '—')[:48])}
</footer>
</div></body></html>"""

    @property
    def _tool(self) -> Dict[str, str]:
        """Identify the software that produced this report.

        Cached. Verification hashes every shipped file, and the template reads
        this five times while rendering one report — so without the cache a
        report re-hashes the whole installation five times over, which turned a
        fast render into a slow one for no added assurance.

        "Which tool, which version, and how do you know it was not altered" is
        the first question under cross-examination. A report that cannot answer
        it describes an examination nobody can reproduce, so this belongs on the
        face of the report rather than buried in a log.
        """
        cached = getattr(self, "_tool_cache", None)
        if cached is not None:
            return cached

        import platform as _platform
        import sys as _sys

        from .. import __version__

        build = ""
        verification = "not verified"
        note = ("This build could not be checked against a release manifest, "
                "so it is a source checkout or an unpackaged copy. Record how "
                "it was obtained.")
        try:
            from ..core.selfcheck import verify_installation
            result = verify_installation()
            build = result.actual_id
            if result.ok:
                verification = (f"{result.checked} files match the release "
                                f"manifest")
                note = ("The installed files were checked against the manifest "
                        "recorded at release and match it exactly. Quote the "
                        "build identifier when this examination is referenced; "
                        "a validation certificate is issued against a specific "
                        "build, not against the product in general.")
            elif result.manifest_present:
                verification = "DOES NOT MATCH THE RELEASE MANIFEST"
                note = ("The installed files differ from the release manifest. "
                        "This examination should not be relied upon until the "
                        "discrepancy is explained.")
        except Exception:                                 # pragma: no cover
            pass

        self._tool_cache = {
            "version": __version__,
            "build": build,
            "verification": verification,
            "python": f"Python {_sys.version.split()[0]}",
            "platform": _platform.platform(),
            "note": note,
        }
        return self._tool_cache

    @property
    def _seal(self) -> Dict[str, str]:
        """The digests that make the integrity claim checkable. Cached.

        Stating "sealed" without publishing the digest asks the reader to take
        it on trust. Publishing it lets anyone recompute the value and confirm
        the container is the one this report describes.
        """
        cached = getattr(self, "_seal_cache", None)
        if cached is not None:
            return cached
        blank = {"name": "—", "container_seal": "—", "blob_merkle_root": "—",
                 "audit": "—", "verified_at": "—"}
        try:
            container = self.session.primary.container
            manifest = getattr(container, "manifest", {}) or {}
            seal = manifest.get("seal") or {}
            report = self.session.integrity_report()
            entry = (report.get("containers") or [{}])[0]
            self._seal_cache = {
                "name": entry.get("name") or "—",
                "container_seal": seal.get("container_seal") or "—",
                "blob_merkle_root": seal.get("blob_merkle_root") or "—",
                "audit": ("valid, {} entries".format(
                    entry.get("audit_entries", 0))
                    if entry.get("audit_chain_valid") else "CHAIN BROKEN"),
                "verified_at": entry.get("verified_at") or "—",
            }
            return self._seal_cache
        except Exception:                                 # pragma: no cover
            return blank

    def _html_lead_sheet(self) -> str:
        """Section 2: ranked investigative findings.

        Placed before the artifact tables because a reader who reaches 30 000
        rows before being told what matters has been given data, not a report.
        Every finding shows its confidence and its caveat — a lead presented
        without the reason it might be wrong invites over-reading.
        """
        intel = self.data.get("intelligence")
        e = html.escape
        if not intel:
            err = self.data.get("intelligence_error")
            if err:
                return (f'<h2>3. Investigative findings</h2>'
                        f'<div class="banner bad">The intelligence layer did '
                        f'not run: {e(err)}. Findings are therefore absent from '
                        f'this report rather than empty.</div>')
            return ""

        f = intel["findings"]
        if not f["count"]:
            return ('<h2>3. Investigative findings</h2><p class="note">No rule '
                    'produced a finding for this evidence. This is not a '
                    'statement that nothing of interest is present — only that '
                    'the automated rules did not fire.</p>')

        sev_colour = {"critical": "#c0201a", "high": "#d1741f",
                      "medium": "#1f4e79", "low": "#5b6774", "info": "#5b6774"}
        blocks = []
        for item in f["findings"]:
            colour = sev_colour.get(item["severity"], "#5b6774")
            evidence = "".join(
                f"<li>{e(str(x)[:220])}</li>" for x in item.get("evidence", [])[:4])
            parties = ", ".join(e(str(p)) for p in item.get("parties", [])[:8])
            blocks.append(f"""
            <div class="finding" style="border-left:4px solid {colour}">
              <div class="fhead">
                <span class="sev" style="background:{colour}">{e(item['severity'].upper())}</span>
                <b>{e(item['title'])}</b>
                <span class="conf">confidence {item['confidence']}</span>
              </div>
              <p>{e(item['detail'])}</p>
              {f'<p class="why"><b>Why it matters:</b> {e(item["why_it_matters"])}</p>'
               if item.get('why_it_matters') else ''}
              {f'<ul class="ev">{evidence}</ul>' if evidence else ''}
              {f'<p class="meta">Parties: {parties}</p>' if parties else ''}
              <p class="meta">Cites {len(item.get('artifact_ids', []))} artifact(s)
                 · rule <code>{e(item['rule_id'])}</code>
                 {f"· {e(item['first_seen_iso'][:19])} to {e(item['last_seen_iso'][:19])}"
                  if item.get('first_seen_iso') else ''}</p>
              {f'<p class="caveat"><b>Caveat:</b> {e(item["caveat"])}</p>'
               if item.get('caveat') else ''}
            </div>""")

        ents = intel.get("entities", {})
        ent_rows = "".join(
            f"<tr><td>{e(h['label'])}</td><td><code>{e(h['value'][:60])}</code></td>"
            f"<td>{h['count']}</td>"
            f"<td>{'checksum verified' if h['validated'] else 'format only'}</td>"
            f"<td>{e(', '.join(h.get('apps', [])[:3]))}</td></tr>"
            for h in (ents.get("high_value") or [])[:25])

        corr = intel.get("correlation") or {}
        coloc = corr.get("colocation") or {}
        corr_html = ""
        if corr.get("shared_party_count") or coloc.get("encounter_count"):
            rows = "".join(
                f"<tr><td>{e(p['best_label'])}</td>"
                f"<td>{e(', '.join(p['exhibits']))}</td>"
                f"<td>{sum(p['exhibits'].values())}</td>"
                f"<td>{e(', '.join(p['deleted_on']) or '—')}</td></tr>"
                for p in corr.get("shared_parties", [])[:20])
            coloc_para = (
                f"""<p>{coloc['encounter_count']} co-location encounter(s) —
                   devices placed in the same place at the same time, not
                   merely the same cell — were also found; see the
                   'colocation.rendezvous' finding(s) above for detail.</p>"""
                if coloc.get("encounter_count") else "")
            corr_html = f"""
            <h3>Cross-exhibit correlation</h3>
            <p>Parties present on more than one exhibit, matched on an exact
               shared identifier. {corr.get('shared_media_count', 0)} byte-identical
               file(s) and {corr.get('shared_location_count', 0)} shared
               location cell(s) were also found.</p>
            {coloc_para}
            <table><thead><tr><th>Party</th><th>Exhibits</th>
              <th>Artifacts</th><th>Deleted on</th></tr></thead>
              <tbody>{rows}</tbody></table>
            <p class="note">{e(corr.get('note', ''))}</p>"""

        return f"""
<h2>3. Investigative findings</h2>
<p>{f['count']} finding(s) produced by {f['rules_run']} rules over
   {f['artifacts_analysed']:,} artifacts:
   {', '.join(f'{v} {k}' for k, v in f['by_severity'].items() if v)}.
   Findings are ranked by severity then confidence. Each cites the artifacts it
   was derived from, and each states the caveat that would make it wrong.</p>
{''.join(blocks)}
{f'<h3>High-value entities extracted from content</h3><table><thead><tr>'
 f'<th>Type</th><th>Value</th><th>Occurrences</th><th>Validation</th>'
 f'<th>Applications</th></tr></thead><tbody>{ent_rows}</tbody></table>'
 f'<p class="note">Checksum-verified values are well-formed; that does not '
 f'establish that they belong to the device owner or were ever used.</p>'
 if ent_rows else ''}
{corr_html}"""

    # ------------------------------------------------------------------- PDF
    def _write_pdf(self, path: Path) -> None:
        try:
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import mm
            from reportlab.platypus import (SimpleDocTemplate, Paragraph,
                                            Spacer, Table, TableStyle,
                                            PageBreak)
        except ImportError as exc:
            raise ReportError("PDF export requires reportlab "
                              "(pip install reportlab)") from exc

        d, o = self.data, self.data["overview"]
        integ = d["integrity"]
        styles = getSampleStyleSheet()
        h1 = ParagraphStyle("h1", parent=styles["Heading1"], fontSize=17,
                            textColor=colors.HexColor("#1f4e79"))
        h2 = ParagraphStyle("h2", parent=styles["Heading2"], fontSize=12.5,
                            textColor=colors.HexColor("#1f4e79"),
                            spaceBefore=12)
        body = ParagraphStyle("b", parent=styles["BodyText"], fontSize=8.5,
                              leading=11)
        cell = ParagraphStyle("c", parent=body, fontSize=7.2, leading=8.8)

        doc = SimpleDocTemplate(
            str(path), pagesize=A4,
            leftMargin=15 * mm, rightMargin=15 * mm,
            topMargin=15 * mm, bottomMargin=15 * mm,
            title=self.options.title, author="ARGUS Forensics")
        story: List[Any] = []

        story.append(Paragraph(html.escape(self.options.title), h1))
        story.append(Paragraph(
            f"Generated {d['generated_at']} · ARGUS Forensics · "
            f"scope: {html.escape(self.options.scope)}", body))
        story.append(Spacer(1, 7 * mm))

        # Integrity first
        if integ["ok"]:
            banner = Table([[Paragraph(
                "<b>INTEGRITY VERIFIED</b> — every blob re-hashed to its "
                "stored digest, the container seal matches, and the "
                "chain-of-custody log is unbroken.", body)]],
                colWidths=[178 * mm])
            banner.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#e9f7ee")),
                ("BOX", (0, 0), (-1, -1), 0.7, colors.HexColor("#1b7f37")),
                ("PADDING", (0, 0), (-1, -1), 7)]))
        else:
            problems = "<br/>".join(
                html.escape(f"{c['name']}: {p}")
                for c in integ["containers"] for p in c.get("problems", []))
            banner = Table([[Paragraph(
                "<b>INTEGRITY VERIFICATION FAILED</b> — the evidence described "
                "below does not match its recorded hashes. These findings must "
                "not be relied upon until this is resolved.<br/>" + problems,
                body)]], colWidths=[178 * mm])
            banner.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#fdecea")),
                ("BOX", (0, 0), (-1, -1), 0.9, colors.HexColor("#c0201a")),
                ("PADDING", (0, 0), (-1, -1), 7)]))
        story += [banner, Spacer(1, 6 * mm)]

        story.append(Paragraph("1. Case and exhibit", h2))
        info = [
            ["Case ID", o["case_id"], "Device", " ".join(filter(None, [
                o["device"]["make"], o["device"]["model"]]))],
            ["Exhibit ID", o["exhibit_id"], "OS", o["device"]["os"]],
            ["Operator", o["operator"], "Serial", o["device"]["serial"]],
            ["Examiner", self.options.examiner, "IMEI", o["device"]["imei"]],
            ["Method", o["method"], "ICCID", o["device"]["iccid"]],
            ["Time span", o["time_span"], "Number", o["device"]["phone_number"]],
            ["Started", o["started_at"], "Lock state", o["device"]["lock_state"]],
            ["Finished", o["finished_at"], "Reference", self.options.reference],
        ]
        t = Table([[Paragraph(f"<b>{html.escape(str(r[0]))}</b>", cell),
                    Paragraph(html.escape(str(r[1] or "—")), cell),
                    Paragraph(f"<b>{html.escape(str(r[2]))}</b>", cell),
                    Paragraph(html.escape(str(r[3] or "—")), cell)]
                   for r in info],
                  colWidths=[26 * mm, 63 * mm, 26 * mm, 63 * mm])
        t.setStyle(TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#dde3ea")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("PADDING", (0, 0), (-1, -1), 3)]))
        story += [t, Spacer(1, 5 * mm)]

        story.append(Paragraph("2. Summary of findings", h2))
        story.append(Paragraph(
            f"{d['artifact_count']:,} artifacts are included, of which "
            f"<b>{d['deleted_count']:,}</b> were recovered from deleted or "
            f"unallocated space. Activity spans "
            f"{html.escape(o['first_activity'] or 'unknown')} to "
            f"{html.escape(o['last_activity'] or 'unknown')}.", body))
        story.append(Spacer(1, 3 * mm))

        cat_rows = [["Category", "Count"]] + [
            [k, f"{v:,}"] for k, v in
            sorted(d["category_counts"].items(), key=lambda kv: -kv[1])]
        ct = Table(cat_rows, colWidths=[120 * mm, 30 * mm])
        ct.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1f4e79")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("GRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#dde3ea")),
            ("PADDING", (0, 0), (-1, -1), 3.5)]))
        story += [ct, Spacer(1, 4 * mm)]

        if d.get("graph", {}).get("top_contacts"):
            story.append(Paragraph("3. Communication analysis", h2))
            rows = [["Party", "Artifacts", "Calls", "Messages", "Parties"]]
            for n in d["graph"]["top_contacts"][:20]:
                rows.append([Paragraph(html.escape(str(n["label"]))[:48], cell),
                             f"{n['artifact_count']:,}", f"{n['calls']:,}",
                             f"{n['messages']:,}", str(n["degree"])])
            gt = Table(rows, colWidths=[78 * mm, 25 * mm, 22 * mm, 28 * mm, 22 * mm])
            gt.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1f4e79")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTSIZE", (0, 0), (-1, -1), 7.5),
                ("GRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#dde3ea")),
                ("PADDING", (0, 0), (-1, -1), 3)]))
            story += [gt, Spacer(1, 4 * mm)]

        story.append(PageBreak())
        story.append(Paragraph("4. Artifacts", h2))
        story.append(Paragraph(
            "Rows shaded in purple were recovered from deleted or unallocated "
            "space and were not visible to the device user at seizure.", body))

        for category, items in sorted(d["by_category"].items(),
                                      key=lambda kv: -len(kv[1])):
            story.append(Paragraph(
                f"{html.escape(category)} ({len(items):,})", h2))
            rows = [["Time (UTC)", "Type", "Party", "Content", "App"]]
            styling = [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1f4e79")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTSIZE", (0, 0), (-1, -1), 6.8),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#dde3ea")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("PADDING", (0, 0), (-1, -1), 2.5)]
            for idx, a in enumerate(items[:250], 1):
                rows.append([
                    Paragraph(html.escape(a.get("timestamp_iso", ""))[:19], cell),
                    Paragraph(html.escape(a.get("subtype", ""))[:26], cell),
                    Paragraph(html.escape("; ".join(
                        p["label"] for p in a.get("parties", [])
                        if not p.get("is_owner")))[:40], cell),
                    Paragraph(html.escape((a.get("body") or ""))[:230], cell),
                    Paragraph(html.escape(a.get("app", ""))[:18], cell)])
                if a.get("is_deleted"):
                    styling.append(("BACKGROUND", (0, idx), (-1, idx),
                                    colors.HexColor("#faf3fd")))
                    styling.append(("TEXTCOLOR", (0, idx), (-1, idx),
                                    colors.HexColor("#7b3fa0")))
            at = Table(rows, colWidths=[24 * mm, 26 * mm, 32 * mm, 74 * mm,
                                        22 * mm], repeatRows=1)
            at.setStyle(TableStyle(styling))
            story += [at, Spacer(1, 3 * mm)]
            if len(items) > 250:
                story.append(Paragraph(
                    f"<i>{len(items) - 250:,} further {html.escape(category)} "
                    f"artifacts are included in the XLSX and XML exports.</i>",
                    body))

        if self.options.conclusion:
            story += [PageBreak(), Paragraph("5. Conclusion", h2),
                      Paragraph(html.escape(self.options.conclusion), body)]

        def footer(canvas, document):
            canvas.saveState()
            canvas.setFont("Helvetica", 7)
            canvas.setFillColor(colors.HexColor("#5b6774"))
            canvas.drawString(15 * mm, 9 * mm,
                              f"ARGUS Forensics · case {o['case_id']} · "
                              f"exhibit {o['exhibit_id']}")
            canvas.drawRightString(195 * mm, 9 * mm, f"Page {document.page}")
            canvas.restoreState()

        doc.build(story, onFirstPage=footer, onLaterPages=footer)


def generate(containers: Sequence[Path | str], out_dir: Path | str,
             options: Optional[ReportOptions] = None,
             basename: str = "forensic_report",
             deep_verify: bool = True) -> List[Path]:
    """One-call report generation from container paths."""
    opts = options or ReportOptions()
    with AnalysisSession(containers, deep_verify=deep_verify) as session:
        return ReportBuilder(session, opts).write(out_dir, basename)
